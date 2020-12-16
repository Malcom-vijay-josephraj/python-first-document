from docx import Document   #pip3 install python-docx or pip3 install python docx
from docx.shared import Inches
import pyttsx3              #pip3 install pyttsx3

def speak(text):
     pyttsx3.speak(text)
     
     
document = Document()

#profile picture
document.add_picture('malcom.jpg' , width = Inches(3.0))
#details
name = input('What is your name?')
speak ('hello ' + name + 'how are you?')
speak('What is your phone number?')
phone_number = input('What is your phone number?')
email = input('What is your email id?')

document.add_paragraph(
     name + ' | ' + phone_number + ' | ' + email)

#about me
document.add_heading('About me')
document.add_paragraph(input('Tell about yourself ? '))


#work experience
document.add_heading('Experience :')
m = document.add_paragraph()

company = input('Enter the comapny name')
date = input('From date')
date2 = input('To date')

m.add_run(company + ' ').bold = True
m.add_run(date + ' - ' + date2 + '\n').italic = True

while True:
     more_experience =input('Do you have more experience ? Yes or No ')
     if more_experience.lower() == 'yes' :
          m = document.add_paragraph()

          company = input('Enter the comapny name' )
          date = input('From date')
          date2 = input('To date')

          m.add_run(company +  '  ').bold = True
          m.add_run(date + ' - ' + date2 + '\n').italic = True

     else :
          break



#skills
document.add_heading('SKILLS :')
skill = input('Enter your skill :')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True :
     has_more_skills = input('Enter extra skill ? Yes or No')
     if has_more_skills.lower() == 'yes' :
          skill = input('Enter your skill :')
          p = document.add_paragraph(skill)
          p.style = 'List Bullet'
     else :
          break


#FOOTER
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "Cv is generated using Malcom and with help of python"











document.save('cv.docx')